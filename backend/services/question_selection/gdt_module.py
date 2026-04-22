# -*- coding: utf-8 -*-
"""
GDT Module — Graph / Data-structure / Table rendering utilities.

Provides pure functions for generating visual assets from structured data.
Import and call these functions; do not execute this file directly.
"""

import ast
import os
import re
import matplotlib
matplotlib.use("Agg")  # non-interactive backend, safe for server use
import matplotlib.pyplot as plt

OUTPUT_DIR = "gdt_outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def _save_figure(path: str) -> str:
    plt.savefig(path, bbox_inches="tight", pad_inches=0.05, dpi=150)
    plt.close()
    return path


def _safe_parse_edge_labels(raw: dict) -> dict:
    """Convert string-keyed edge_labels dict to tuple-keyed dict safely."""
    result = {}
    for k, v in raw.items():
        try:
            parsed = ast.literal_eval(k)
            if isinstance(parsed, tuple):
                result[parsed] = v
        except (ValueError, SyntaxError):
            pass
    return result


# ─────────────────────────────────────────────────────────────
# VISUAL GENERATORS
# ─────────────────────────────────────────────────────────────

def generate_plot(content: dict, filename: str) -> str:
    """Render a line plot from {x, y, xlabel, ylabel, title} and save as PNG."""
    plt.figure(figsize=(3.5, 2.5))
    plt.plot(content["x"], content["y"], marker="o", color="black")
    plt.xlabel(content.get("xlabel", ""), fontsize=8)
    plt.ylabel(content.get("ylabel", ""), fontsize=8)
    plt.title(content.get("title", ""), fontsize=9)
    plt.grid(True, linestyle="--", linewidth=0.5)
    return _save_figure(os.path.join(OUTPUT_DIR, filename))


def generate_graph_ds(content: dict, filename: str) -> str:
    """Render a NetworkX graph from edges/edge_labels and save as PNG."""
    import networkx as nx

    G = nx.DiGraph() if content.get("directed", False) else nx.Graph()
    G.add_edges_from(content["edges"])

    plt.figure(figsize=(3, 2.5))
    pos = nx.spring_layout(G, seed=42)
    nx.draw(G, pos, with_labels=True, node_size=800, node_color="white",
            edge_color="black", font_size=8, edgecolors="black")

    if "edge_labels" in content:
        edge_labels = _safe_parse_edge_labels(content["edge_labels"])
        if edge_labels:
            nx.draw_networkx_edge_labels(G, pos, edge_labels=edge_labels, font_size=7)

    path = os.path.join(OUTPUT_DIR, filename)
    plt.savefig(path, bbox_inches="tight", pad_inches=0.05, dpi=150)
    plt.close()
    return path


def generate_formula(formula_latex: str, filename: str) -> str:
    """Render a LaTeX formula string as a small PNG image."""
    fig, ax = plt.subplots(figsize=(2, 0.6))
    ax.text(0.5, 0.5, f"${formula_latex}$", fontsize=12, ha="center", va="center")
    ax.axis("off")
    path = os.path.join(OUTPUT_DIR, filename)
    plt.savefig(path, bbox_inches="tight", pad_inches=0.0, dpi=200)
    plt.close()
    return path


# ─────────────────────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────────────────────

def add_table_with_border(doc, content: dict) -> None:
    """Add a bordered table to a python-docx Document."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table = doc.add_table(rows=1, cols=len(content["headers"]))
    for i, h in enumerate(content["headers"]):
        table.rows[0].cells[i].text = str(h)

    for row in content["rows"]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                elem = OxmlElement(f"w:{edge}")
                elem.set(qn("w:val"), "single")
                elem.set(qn("w:sz"), "4")
                elem.set(qn("w:color"), "000000")
                borders.append(elem)
            tcPr.append(borders)


def add_text_with_formula_block(doc, text: str) -> None:
    """Split text on $...$ formula markers and add each part to a Document."""
    from docx.shared import Inches

    parts = re.split(r"(\$.*?\$)", text)
    for part in parts:
        if part.startswith("$") and part.endswith("$"):
            formula = part.strip("$")
            img_path = generate_formula(formula, "_temp_formula.png")
            doc.add_picture(img_path, width=Inches(2))
        else:
            if part.strip():
                p = doc.add_paragraph(part)
                p.paragraph_format.space_after = 0
                p.paragraph_format.space_before = 0


def render_gdt_to_docx(doc, gdt_blocks: list, question_id: str) -> None:
    """Append all GDT blocks (table / plot / graph_ds) to a Document."""
    from docx.shared import Inches

    for idx, g in enumerate(gdt_blocks):
        gtype = g.get("type", "")
        content = g.get("content", {})

        if gtype == "table":
            add_table_with_border(doc, content)

        elif gtype == "plot":
            img = generate_plot(content, f"{question_id}_plot_{idx}.png")
            doc.add_picture(img, width=Inches(3.5))

        elif gtype == "graph_ds":
            img = generate_graph_ds(content, f"{question_id}_graph_{idx}.png")
            doc.add_picture(img, width=Inches(3.5))

        elif gtype == "formula":
            img = generate_formula(content, f"{question_id}_formula_{idx}.png")
            doc.add_picture(img, width=Inches(2))


# ─────────────────────────────────────────────────────────────
# PDF HELPERS
# ─────────────────────────────────────────────────────────────

def parse_text_with_formula_pdf(text: str, elements: list) -> None:
    """Split text on $...$ markers and append Paragraph / Image elements."""
    from reportlab.platypus import Paragraph, Image
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch

    style = getSampleStyleSheet()["Normal"]
    parts = re.split(r"(\$.*?\$)", text)
    for part in parts:
        if part.startswith("$") and part.endswith("$"):
            formula = part.strip("$")
            img_path = generate_formula(formula, "_temp_formula.png")
            elements.append(Image(img_path, width=2 * inch, height=0.4 * inch))
        else:
            if part.strip():
                elements.append(Paragraph(part, style))


def create_pdf_table(content: dict):
    """Return a ReportLab Table with borders from {headers, rows}."""
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors

    data = [content["headers"]] + [list(map(str, r)) for r in content["rows"]]
    table = Table(data)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    return table


def render_gdt_to_pdf(gdt_blocks: list, elements: list, question_id: str) -> None:
    """Append all GDT blocks (table / plot / graph_ds / formula) as PDF flowables."""
    from reportlab.platypus import Image, Spacer
    from reportlab.lib.units import inch

    for idx, g in enumerate(gdt_blocks):
        gtype = g.get("type", "")
        content = g.get("content", {})
        elements.append(Spacer(1, 4))

        if gtype == "table":
            elements.append(create_pdf_table(content))

        elif gtype == "plot":
            img_path = generate_plot(content, f"{question_id}_plot_{idx}.png")
            elements.append(Image(img_path, width=3.5 * inch, height=2.5 * inch))

        elif gtype == "graph_ds":
            img_path = generate_graph_ds(content, f"{question_id}_graph_{idx}.png")
            elements.append(Image(img_path, width=3.5 * inch, height=2.5 * inch))

        elif gtype == "formula":
            img_path = generate_formula(content, f"{question_id}_formula_{idx}.png")
            elements.append(Image(img_path, width=2 * inch, height=0.4 * inch))
