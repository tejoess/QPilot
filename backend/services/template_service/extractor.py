"""
backend/services/template_service/extractor.py

Extracts placeholder tokens and question pattern from a DOCX template.

Key conventions in the ACPCE IA Template:
- Global placeholders: [subject], [class], [marks], [date], [duration], [exam_name]
- Per-row question placeholders: [1a], [1b], [2a], [2b], ... (question text cell)
- Per-row cognitive level placeholder: [cl] - in THE SAME ROW as the question

Pattern is inferred from which question IDs exist:
  e.g., [1a],[1b],[1c],[1d],[1e] → Section 1: 5 questions
        [2a],[2b] through [6a],[6b] → Sections 2-6: 2 questions each
Marks are read from cells containing patterns like "[2]" or "[5]" (numeric-only brackets).
"""

import re
from typing import List, Dict, Any
from docx import Document


# Placeholders that are NOT question IDs
GLOBAL_PLACEHOLDERS = {
    "[subject]", "[class]", "[marks]", "[date]",
    "[duration]", "[exam_name]", "[cl]",
}

# Regex to match bracket tokens: [sometext]
BRACKET_RE = re.compile(r'\[([^\[\]]+)\]')

# Regex for question id: 1-2 digits followed by a letter(s) — e.g. [1a], [2b], [10a]
QUESTION_ID_RE = re.compile(r'^(\d{1,2})([a-z]+)$', re.IGNORECASE)

# Regex for a marks cell: only a number inside brackets, like [2] or [05]
MARKS_RE = re.compile(r'^\d+$')


def _iter_all_text_in_doc(doc: Document):
    """Yield (text, paragraph_or_cell) for all paragraphs and table cells."""
    for para in doc.paragraphs:
        yield para.text, para

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para.text, para


def extract_placeholders(docx_path: str) -> Dict[str, Any]:
    """
    Extract all [placeholder] tokens from a DOCX template.

    Returns:
        {
            "global": ["[subject]", "[class]", ...],
            "question_ids": ["[1a]", "[1b]", "[2a]", ...],
            "marks_cells": ["[2]", "[5]", ...],
            "all": [<all tokens found>]
        }
    """
    doc = Document(docx_path)
    all_tokens = set()

    for text, _ in _iter_all_text_in_doc(doc):
        for match in BRACKET_RE.finditer(text):
            token = f"[{match.group(1)}]"
            all_tokens.add(token)

    global_ph = []
    question_ids = []
    marks_cells = []

    for token in sorted(all_tokens):
        inner = token[1:-1]  # strip []
        if token in GLOBAL_PLACEHOLDERS:
            global_ph.append(token)
        elif QUESTION_ID_RE.match(inner):
            question_ids.append(token)
        elif MARKS_RE.match(inner):
            marks_cells.append(token)

    # Sort question IDs naturally: [1a] < [1b] < [2a] ...
    question_ids.sort(key=lambda t: (
        int(QUESTION_ID_RE.match(t[1:-1]).group(1)),
        QUESTION_ID_RE.match(t[1:-1]).group(2)
    ))

    return {
        "global": global_ph,
        "question_ids": question_ids,
        "marks_cells": marks_cells,
        "all": sorted(all_tokens),
    }


def extract_pattern(docx_path: str) -> List[Dict]:
    """
    Infer the question paper structural pattern from template placeholders.

    Groups question IDs by their numeric prefix, then tries to find the marks
    for that question group from nearby numeric-bracket cells in the same table row.

    Returns a list of section dicts:
        [
            {"question_num": 1, "parts": ["a","b","c","d","e"], "num_parts": 5, "marks_per_part": 2},
            {"question_num": 2, "parts": ["a","b"], "num_parts": 2, "marks_per_part": 5},
            ...
        ]
    """
    doc = Document(docx_path)

    # Build: question_num -> {parts: [letter...], marks_per_part: int}
    sections: Dict[int, Dict] = {}

    def _ensure_section(qnum: int):
        if qnum not in sections:
            sections[qnum] = {"parts": [], "marks_per_part": 0}

    # Pass 1: detect question IDs from all text blocks (paragraphs + table-cell paragraphs).
    # This enables pattern extraction even when placeholders are not in tables.
    for text, _ in _iter_all_text_in_doc(doc):
        if not text:
            continue

        found_qids = []
        local_marks = []

        for match in BRACKET_RE.finditer(text):
            inner = match.group(1).strip()
            qid_match = QUESTION_ID_RE.match(inner)
            if qid_match:
                qnum = int(qid_match.group(1))
                part = qid_match.group(2).lower()
                found_qids.append((qnum, part))
            elif MARKS_RE.match(inner):
                local_marks.append(int(inner))

        if not found_qids:
            continue

        for qnum, part in found_qids:
            _ensure_section(qnum)
            if part not in sections[qnum]["parts"]:
                sections[qnum]["parts"].append(part)

        # If a paragraph contains both a question id and numeric mark, use it as fallback.
        if local_marks:
            for qnum, _ in found_qids:
                if sections[qnum]["marks_per_part"] == 0:
                    sections[qnum]["marks_per_part"] = local_marks[0]

    # Pass 2: preserve row-aware marks inference for table templates where marks appear
    # in a neighboring cell of the same row.
    for table in doc.tables:
        for row in table.rows:
            cell_texts = [" ".join(p.text.strip() for p in cell.paragraphs) for cell in row.cells]
            row_full_text = " ".join(cell_texts)

            found_qids = []
            for match in BRACKET_RE.finditer(row_full_text):
                inner = match.group(1).strip()
                qid_match = QUESTION_ID_RE.match(inner)
                if qid_match:
                    qnum = int(qid_match.group(1))
                    part = qid_match.group(2).lower()
                    found_qids.append((qnum, part))

            row_marks = []
            for cell_text in cell_texts:
                for match in BRACKET_RE.finditer(cell_text):
                    inner = match.group(1).strip()
                    if MARKS_RE.match(inner):
                        row_marks.append(int(inner))

            for qnum, part in found_qids:
                _ensure_section(qnum)
                if part not in sections[qnum]["parts"]:
                    sections[qnum]["parts"].append(part)

                if row_marks and sections[qnum]["marks_per_part"] == 0:
                    sections[qnum]["marks_per_part"] = row_marks[0]

    # Build ordered pattern list
    pattern = []
    for qnum in sorted(sections.keys()):
        sec = sections[qnum]
        parts = sorted(sec["parts"])
        pattern.append({
            "question_num": qnum,
            "parts": parts,
            "num_parts": len(parts),
            "marks_per_part": sec["marks_per_part"],
        })

    return pattern


def pattern_matches_paper(template_pattern: List[Dict], paper_sections: List[Dict]) -> bool:
    """
    Check if the generated paper's sections match the template pattern.

    template_pattern: output of extract_pattern()
    paper_sections: list of {"num_questions": N, "marks_per_question": M}
        (the same format as what the frontend sends as paper_pattern sections)

    Returns True only if:
    - same number of questions per group
    - same marks per question for each group
    """
    if not template_pattern or not paper_sections:
        return False

    if len(template_pattern) != len(paper_sections):
        return False

    for tpl, gen in zip(template_pattern, paper_sections):
        tpl_n = tpl.get("num_parts", 0)
        tpl_m = tpl.get("marks_per_part", 0)
        gen_n = gen.get("num_questions", gen.get("numQuestions", 0))
        gen_m = gen.get("marks_per_question", gen.get("marksPerQuestion", 0))

        if tpl_n != gen_n or tpl_m != gen_m:
            return False

    return True
