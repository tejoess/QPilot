"""
backend/services/template_service/renderer.py

Renders a generated paper JSON into a DOCX template by replacing placeholders.

Placeholder conventions:
  Global (anywhere in doc):
    [subject]    → exam subject name
    [class]      → class / grade
    [marks]      → total marks
    [date]       → exam date
    [duration]   → exam duration
    [exam_name]  → examination name / title

  Per-row (inside table rows):
    [1a], [1b], [2a], [2b] ... → question text
    [cl]  → cognitive level (in same row as the question)
             mapped from Bloom taxonomy level via bloom_to_cl()

  Optional marks cells:
    [2], [5] etc. — these are left as-is (they define the template structure)

Bloom → CL mapping (ACPCE standard):
    remember / understand → L1 (Knowledge/Comprehension)
    apply / analyze       → L2 (Application/Analysis)
    evaluate / create     → L3 (Synthesis/Evaluation)
"""

import re
import copy
import os
import tempfile
from typing import Dict, List, Any, Optional
from docx import Document
from docx.oxml.ns import qn

try:
    from backend.services.question_selection.gdt_module import (
        add_text_with_formula_block,
        render_gdt_to_docx,
    )
    _GDT_AVAILABLE = True
except Exception:
    _GDT_AVAILABLE = False


BLOOM_TO_CL: Dict[str, str] = {
    "remember":   "L1",
    "understand": "L1",
    "apply":      "L2",
    "analyze":    "L2",
    "analyse":    "L2",
    "evaluate":   "L3",
    "create":     "L3",
}

BRACKET_RE = re.compile(r'\[([^\[\]]+)\]')
QUESTION_ID_RE = re.compile(r'^(\d{1,2})([a-z]+)$', re.IGNORECASE)


def bloom_to_cl(level: str) -> str:
    """Map a Bloom taxonomy level string to a CL abbreviation."""
    return BLOOM_TO_CL.get(str(level).lower().strip(), "L1")


def _build_question_map(paper_json: Dict) -> Dict[str, Dict]:
    """
    Build a map: question_id_str → {text, bloom, cl}
    from the generated paper JSON.

    Actual paper JSON structure (from final_paper.json):
    {
        "sections": [
            {
                "section_name": "Q1",
                "questions": [
                    {
                        "id": "...",
                        "question_number": "1a",   ← the slot key
                        "question_text": "...",    ← the text
                        "bloom_level": "Remember", ← Title-case
                        ...
                    }
                ]
            }
        ]
    }

    Also handles legacy shapes with id/qid and question/text fields.
    """
    qmap: Dict[str, Dict] = {}

    def _add(q_id: Any, text: str, bloom: str, gdt: list = None):
        cid = str(q_id).lower().strip()
        if not cid:
            return
        qmap[cid] = {
            "text": text or "",
            "bloom": bloom,
            "cl": bloom_to_cl(bloom),
            "gdt": gdt or [],
        }

    sections = paper_json.get("sections", [])
    if sections:
        for sec_idx, section in enumerate(sections):
            questions = section.get("questions", [])
            for q_idx, q in enumerate(questions):
                # Prefer question_number (actual format), fall back to id/qid
                q_id = (
                    q.get("question_number")
                    or q.get("id")
                    or q.get("qid")
                )
                if not q_id:
                    # Synthesise id from section + position: section 0 q 0 → "1a"
                    q_id = f"{sec_idx + 1}{chr(ord('a') + q_idx)}"

                # Prefer question_text (actual format), fall back to question/text
                text = (
                    q.get("question_text")
                    or q.get("question")
                    or q.get("text")
                    or ""
                )
                bloom = (
                    q.get("bloom_level")
                    or q.get("bloom")
                    or q.get("cognitive_level")
                    or "remember"
                )
                _add(q_id, text, bloom, q.get("gdt") or [])
    else:
        # Flat list fallback
        flat = paper_json.get("questions", [])
        for q in flat:
            q_id = (
                q.get("question_number")
                or q.get("id")
                or q.get("qid")
            )
            text = (
                q.get("question_text")
                or q.get("question")
                or q.get("text")
                or ""
            )
            bloom = (
                q.get("bloom_level")
                or q.get("bloom")
                or q.get("cognitive_level")
                or "remember"
            )
            if q_id:
                _add(q_id, text, bloom, q.get("gdt") or [])

    print(f"[renderer] Built qmap with {len(qmap)} questions: {list(qmap.keys())}")
    return qmap


def _replace_in_run_text(text: str, replacements: Dict[str, str]) -> str:
    """Replace all [placeholder] tokens in a string with given replacements."""
    lowered = {k.lower(): v for k, v in replacements.items()}
    def _sub(m):
        token = f"[{m.group(1)}]"
        return lowered.get(token.lower(), token)
    return BRACKET_RE.sub(_sub, text)


def _replace_paragraph_placeholders(para, replacements: Dict[str, str]):
    """
    Replace placeholder tokens in a paragraph, preserving run formatting.

    Uses lxml to collect ALL <w:t> text nodes (handles runs inside hyperlinks,
    tracked-change insertions <w:ins>, SDTs, etc. that para.runs misses).
    Reconstructs the full paragraph text, does the replacement, then writes
    back: the first <w:t> gets the full new text; the rest are cleared.
    """
    wt_elements = para._p.findall('.//' + qn('w:t'))
    if not wt_elements:
        return

    full_text = ''.join(t.text or '' for t in wt_elements)
    if not BRACKET_RE.search(full_text):
        return

    new_text = _replace_in_run_text(full_text, replacements)
    if new_text == full_text:
        return

    # Write back: put everything in the first w:t, blank the rest
    wt_elements[0].text = new_text
    for wt in wt_elements[1:]:
        wt.text = ''


def _insert_nested_table(cell, headers: list, rows: list) -> None:
    """
    Insert a properly bordered table directly inside a DOCX table cell using
    lxml XML manipulation (python-docx has no cell.add_table() API).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _border_elem(tag: str) -> Any:
        el = OxmlElement(tag)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        return el

    def _make_cell_xml(value: str, bold: bool = False):
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            tcBorders.append(_border_elem(f"w:{edge}"))
        tcPr.append(tcBorders)
        tc.append(tcPr)
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        if bold:
            rPr = OxmlElement("w:rPr")
            rPr.append(OxmlElement("w:b"))
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = value
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    # Build <w:tbl>
    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tblBorders.append(_border_elem(f"w:{edge}"))
    tblPr.append(tblBorders)
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tbl.append(tblPr)

    tblGrid = OxmlElement("w:tblGrid")
    for _ in headers:
        gridCol = OxmlElement("w:gridCol")
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    # Header row
    tr = OxmlElement("w:tr")
    for h in headers:
        tr.append(_make_cell_xml(str(h), bold=True))
    tbl.append(tr)

    # Data rows
    for row_data in rows:
        tr = OxmlElement("w:tr")
        for val in row_data:
            tr.append(_make_cell_xml(str(val)))
        tbl.append(tr)

    # Append the nested table into the parent cell's XML
    cell._tc.append(tbl)


def _add_gdt_inline(cell, gdt_blocks: list, q_id: str) -> None:
    """
    Append GDT content (table / plot / graph_ds / formula) directly into a
    python-docx table cell, immediately after the question text.
    """
    if not _GDT_AVAILABLE or not gdt_blocks:
        return

    from docx.shared import Inches
    try:
        from backend.services.question_selection.gdt_module import (
            generate_plot, generate_graph_ds, generate_formula,
        )
    except Exception:
        return

    for idx, block in enumerate(gdt_blocks):
        btype = block.get("type", "")
        content = block.get("content", {})

        if btype == "table" and isinstance(content, dict):
            headers = content.get("headers", [])
            rows    = content.get("rows", [])
            _insert_nested_table(cell, headers, rows)

        elif btype == "plot":
            try:
                img_path = generate_plot(content, f"{q_id}_plot_{idx}.png")
                p = cell.add_paragraph()
                p.add_run().add_picture(img_path, width=Inches(3))
            except Exception as e:
                print(f"[renderer] GDT plot render failed: {e}")

        elif btype == "graph_ds":
            try:
                img_path = generate_graph_ds(content, f"{q_id}_graph_{idx}.png")
                p = cell.add_paragraph()
                p.add_run().add_picture(img_path, width=Inches(3))
            except Exception as e:
                print(f"[renderer] GDT graph render failed: {e}")

        elif btype == "formula" and isinstance(content, str):
            try:
                img_path = generate_formula(content, f"{q_id}_formula_{idx}.png")
                p = cell.add_paragraph()
                p.add_run().add_picture(img_path, width=Inches(2))
            except Exception as e:
                print(f"[renderer] GDT formula render failed: {e}")


def _process_table(table, global_replacements: Dict[str, str], qmap: Dict[str, Dict]):
    """
    Process a table's rows. For each row:
    1. Find question ID placeholders in the row → look up question text and CL.
    2. Build row-specific replacements (question text + [cl] for that row's question).
    3. Apply all replacements to every cell in the row.
    4. Append any GDT blocks inline into the question cell.
    """
    for row in table.rows:
        # Collect all text in this row to check for question IDs
        row_text = " ".join(
            " ".join(p.text for p in cell.paragraphs)
            for cell in row.cells
        )

        # Find question IDs in this row
        row_qids: List[str] = []
        for m in BRACKET_RE.finditer(row_text):
            inner = m.group(1)
            if QUESTION_ID_RE.match(inner):
                row_qids.append(inner.lower())

        # Build row-level replacements on top of globals
        row_replacements = dict(global_replacements)

        gdt_cell: Any = None   # cell that holds the primary question placeholder
        primary_q_info: Any = None

        if row_qids:
            primary_qid = row_qids[0]
            q_info = qmap.get(primary_qid)

            print(f"[renderer] Row qids={row_qids}, primary={primary_qid}, found={q_info is not None}")

            if q_info:
                row_replacements[f"[{primary_qid}]"] = q_info["text"]
                row_replacements["[cl]"] = q_info["cl"]

                for qid in row_qids[1:]:
                    q2 = qmap.get(qid)
                    if q2:
                        row_replacements[f"[{qid}]"] = q2["text"]

                # Identify which cell has the question placeholder (scan BEFORE replacement)
                if q_info.get("gdt"):
                    for cell in row.cells:
                        cell_raw = " ".join(p.text for p in cell.paragraphs)
                        if f"[{primary_qid}]" in cell_raw.lower():
                            gdt_cell = cell
                            primary_q_info = q_info
                            break
            else:
                print(f"[renderer] WARNING: '{primary_qid}' not found in qmap. Available keys: {list(qmap.keys())[:10]}")

        # Apply replacements to all paragraphs in all cells in this row
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_paragraph_placeholders(para, row_replacements)

        # Append GDT content inline into the question cell (after text replacement)
        if gdt_cell is not None and primary_q_info is not None:
            _add_gdt_inline(gdt_cell, primary_q_info["gdt"], primary_qid)


def render_template(
    docx_path: str,
    paper_json: Dict,
    metadata: Dict[str, str],
    out_path: Optional[str] = None,
) -> str:
    """
    Render a paper JSON into a DOCX template.

    Args:
        docx_path:   Path to the source DOCX template.
        paper_json:  Generated paper JSON (from pipeline result).
        metadata:    Dict with keys: subject, class, marks, date, duration, exam_name.
        out_path:    Output path for rendered DOCX. Auto-generates if None.

    Returns:
        Absolute path to the rendered DOCX file.
    """
    if out_path is None:
        tmp_dir = tempfile.gettempdir()
        out_path = os.path.join(tmp_dir, f"rendered_paper_{os.getpid()}.docx")

    doc = Document(docx_path)
    qmap = _build_question_map(paper_json)

    # Global replacements
    global_replacements: Dict[str, str] = {
        "[subject]":   metadata.get("subject", ""),
        "[class]":     metadata.get("class", metadata.get("grade", "")),
        "[marks]":     str(metadata.get("marks", metadata.get("total_marks", ""))),
        "[date]":      metadata.get("date", ""),
        "[duration]":  metadata.get("duration", ""),
        "[exam_name]": metadata.get("exam_name", metadata.get("title", "")),
        "[cl]":        "L1",   # fallback — overridden per-row when question is matched
    }

    # Make question-id placeholders replaceable anywhere in the document,
    # including non-table paragraphs (not just table rows).
    for qid, qinfo in qmap.items():
        global_replacements[f"[{qid}]"] = qinfo.get("text", "")

    # Process body paragraphs (non-table content like header)
    for para in doc.paragraphs:
        _replace_paragraph_placeholders(para, global_replacements)

    # Process all tables — GDT is now appended inline into each question's cell
    for table in doc.tables:
        _process_table(table, global_replacements, qmap)

    doc.save(out_path)
    return out_path
